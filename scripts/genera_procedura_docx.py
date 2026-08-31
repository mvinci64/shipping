"""Genera docs/PROC-LAB-04_aggiornamento_fine_giornata.docx nello stesso
formato del regolamento esistente (docs/esempi/REG-LAB-01 Regolamento
interno di Laboratorio.pdf): tabella intestazione con codice/revisione/data,
titolo, sezioni numerate, "Documenti collegati", tabella "Presa visione del
personale" con firme, footer con codice/pagina, firma del responsabile in
chiusura. Va rigenerato da qui se il contenuto cambia — non modificare il
.docx a mano per la struttura (la copia in docs/ resta comunque il master
per le correzioni di testo puntuali, come da CLAUDE.md)."""
import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GRIGIO = RGBColor(0x59, 0x59, 0x59)
CODICE = "PROC-LAB-04"
REVISIONE = "Rev. 00"
RESPONSABILE = "Marcello Vinci"

d = docx.Document()
for s in d.sections:
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)
    s.top_margin = Cm(1.8)
    s.bottom_margin = Cm(1.8)


def shade_cell(cell, colore_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), colore_hex)
    tcPr.append(shd)


def set_cell_text(cell, testo, bold=False, size=9, color=None, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(testo)
    r.bold = bold
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        r.font.color.rgb = color


# --- tabella intestazione (codice / revisione / data) ---
header = d.add_table(rows=1, cols=4)
header.alignment = WD_TABLE_ALIGNMENT.CENTER
header.style = "Table Grid"
set_cell_text(header.cell(0, 0), "VISCOTTA — Procedure interne di Laboratorio", bold=True, size=9)
set_cell_text(header.cell(0, 1), f"Cod. {CODICE}", size=9)
set_cell_text(header.cell(0, 2), REVISIONE, size=9)
set_cell_text(header.cell(0, 3), "Data __/__/2026", size=9)

d.add_paragraph()

# --- titolo ---
titolo = d.add_paragraph()
r = titolo.add_run("Aggiornamento dati di fine giornata")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = NAVY

sub = d.add_paragraph()
r = sub.add_run(
    "Il passaggio che collega la produzione in laboratorio ai dati di lotto e scadenza "
    "sulle etichette dei colli."
)
r.font.size = Pt(11)
r.font.color.rgb = GRIGIO


def par(testo, size=11):
    p = d.add_paragraph()
    r = p.add_run(testo)
    r.font.size = Pt(size)
    return p


def sezione(numero, titolo_sez):
    p = d.add_paragraph()
    r = p.add_run(f"{numero}. {titolo_sez}")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY
    return p


def elenco(voci, numerato=False):
    for i, voce in enumerate(voci, start=1):
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        prefisso = f"{i}. " if numerato else "• "
        r = p.add_run(prefisso + voce)
        r.font.size = Pt(11)


par(
    "La presente procedura integra il Regolamento interno di Laboratorio (REG-LAB-01) e si applica "
    "a chi chiude la giornata di produzione."
)

sezione(1, "Perché questo passaggio")
par(
    "Le etichette dei colli riportano lotto e scadenza del prodotto: un'informazione obbligatoria, "
    "non un dettaglio. Perché sia sempre quella corretta, il sistema deve prima sapere che il lotto "
    "del giorno è stato registrato. Questo passaggio oggi non è automatico: va fatto a mano, una "
    "volta al giorno."
)

sezione(2, "Quando farlo")
par(
    "Una volta al giorno, a fine turno di produzione, dopo che l'ultimo lotto è stato registrato nel "
    "programma di magazzino (EasyFatt) e prima che inizi la stampa delle etichette dei colli. Non serve "
    "farlo più volte al giorno: il lotto è unico per ogni prodotto in ogni giornata, quindi un solo "
    "aggiornamento a fine giornata è sufficiente per tutti gli ordini di quel giorno."
)

sezione(3, "Come farlo, in breve")
elenco([
    "Verificare che l'ultimo lotto del giorno sia stato registrato in EasyFatt.",
    "Lanciare l'aggiornamento verso il sistema BI (il comando già in uso per aggiornare i dati).",
    "Attendere il completamento prima di autorizzare la stampa delle etichette dei colli di quel giorno.",
], numerato=True)

sezione(4, "Cosa succede se questo passaggio viene saltato")
par(
    "Le etichette stampate prima dell'aggiornamento possono riportare il lotto di un giorno precedente, "
    "oppure restare senza lotto. In entrambi i casi la tracciabilità del collo non è garantita: la stampa "
    "va rifatta dopo aver completato l'aggiornamento."
)

sezione(5, "Cosa NON cambia")
elenco([
    "Il modo di produrre e di registrare i lotti in EasyFatt: identico a oggi.",
    "L'aspetto e il formato delle etichette dei colli: cambia solo il momento in cui vanno stampate.",
])

sezione(6, "Responsabile del passaggio")
par(
    "Da assegnare in modo stabile a chi chiude la giornata in laboratorio. Fino alla nomina, il "
    "riferimento per qualsiasi dubbio è Marcello Vinci."
)

d.add_paragraph()
doc_coll = d.add_paragraph()
r = doc_coll.add_run("Documenti collegati")
r.bold = True
r.font.size = Pt(12)
elenco([
    "REG-LAB-01 — Regolamento interno di Laboratorio",
])

d.add_paragraph()
firme_tit = d.add_paragraph()
r = firme_tit.add_run("Presa visione del personale")
r.bold = True
r.font.size = Pt(12)

tabella = d.add_table(rows=8, cols=3)
tabella.style = "Table Grid"
tabella.alignment = WD_TABLE_ALIGNMENT.CENTER
larghezze = [Cm(9.5), Cm(3), Cm(4)]
for i, w in enumerate(larghezze):
    for cella in tabella.columns[i].cells:
        cella.width = w
intestazioni = ["Nome e Cognome", "Data", "Firma"]
for i, testo in enumerate(intestazioni):
    cell = tabella.cell(0, i)
    shade_cell(cell, "1F4E79")
    set_cell_text(cell, testo, bold=True, size=10, white=True)
for riga in range(1, 8):
    for col in range(3):
        tabella.cell(riga, col).text = ""

d.add_page_break()
firma_p = d.add_paragraph()
firma_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = firma_p.add_run("Data ________________")
r.font.size = Pt(11)
r2 = firma_p.add_run("                " + f"{RESPONSABILE} ____________________")
r2.font.size = Pt(11)

# --- footer con codice/revisione/pagina ---
section = d.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.text = f"{CODICE} · {REVISIONE} · Pag. "
run_pag = fp.add_run()
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.set(qn("xml:space"), "preserve")
instrText.text = "PAGE"
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "end")
run_pag._r.append(fldChar1)
run_pag._r.append(instrText)
run_pag._r.append(fldChar2)
fp.add_run("/")
run_pag2 = fp.add_run()
fldChar3 = OxmlElement("w:fldChar")
fldChar3.set(qn("w:fldCharType"), "begin")
instrText2 = OxmlElement("w:instrText")
instrText2.set(qn("xml:space"), "preserve")
instrText2.text = "NUMPAGES"
fldChar4 = OxmlElement("w:fldChar")
fldChar4.set(qn("w:fldCharType"), "end")
run_pag2._r.append(fldChar3)
run_pag2._r.append(instrText2)
run_pag2._r.append(fldChar4)
for run in fp.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = GRIGIO

d.save("docs/PROC-LAB-04_aggiornamento_fine_giornata.docx")
print("salvato docs/PROC-LAB-04_aggiornamento_fine_giornata.docx")
